"""
ResumeXpert – Resume Parser
Extracts raw text, name, email, and phone from PDF / DOCX / Image resumes.

PDF strategy (cascade, stops at first success):
  1. PyPDF2  – fast, handles most modern PDFs
  2. pdfplumber – richer extraction, handles tricky layouts
  3. pdfminer.six – deepest extraction, slowest but most compatible
"""

import re
import os

# Sentinel returned for image files so the caller knows to use Vision AI
IMAGE_FILE_SENTINEL = "__IMAGE_FILE__"


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------
def extract_text(file_path: str) -> str:
    """
    Return raw text from a resume file.
    Supports .pdf, .docx, .jpg, .jpeg, .png.

    For image files returns IMAGE_FILE_SENTINEL — the caller must
    route those to the Gemini Vision endpoint instead.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return _extract_pdf(file_path)
    elif ext == '.docx':
        return _extract_docx(file_path)
    elif ext in ('.jpg', '.jpeg', '.png'):
        print(f"[ResumeParser] Image file detected: {file_path} — routing to Vision AI.")
        return IMAGE_FILE_SENTINEL
    else:
        return ""


# ---------------------------------------------------------------------------
# PDF extraction  — 3-library cascade
# ---------------------------------------------------------------------------
def _extract_pdf(path: str) -> str:
    text = ""

    import gc

    # ── Library 1: PyMuPDF (fitz) ────────────────────────────────────────────
    try:
        import fitz
        text_chunks = []
        with fitz.open(path) as doc:
            for page in doc:
                text_chunks.append(page.get_text())
                del page
        text = "\n".join(text_chunks).strip()
        del text_chunks
        gc.collect()
        
        if text:
            print(f"[ResumeParser] PDF parsed successfully with PyMuPDF ({len(text)} chars).")
            return text
        print("[ResumeParser] PyMuPDF returned empty — trying pdfplumber.")
    except ImportError:
        print("[ResumeParser] PyMuPDF not installed — trying pdfplumber.")
    except Exception as e:
        print(f"[ResumeParser] PyMuPDF failed: {e} — trying pdfplumber.")

    # ── Library 2: pdfplumber ────────────────────────────────────────────────
    try:
        import pdfplumber
        text_chunks = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pg_text = page.extract_text()
                if pg_text:
                    text_chunks.append(pg_text)
                del page
        text = "\n".join(text_chunks).strip()
        del text_chunks
        gc.collect()
        
        if text:
            print(f"[ResumeParser] PDF parsed successfully with pdfplumber ({len(text)} chars).")
            return text
        print("[ResumeParser] pdfplumber returned empty — trying PyPDF2.")
    except ImportError:
        print("[ResumeParser] pdfplumber not installed — trying PyPDF2.")
    except Exception as e:
        print(f"[ResumeParser] pdfplumber failed: {e} — trying PyPDF2.")

    # ── Library 3: PyPDF2 ────────────────────────────────────────────────────
    try:
        import PyPDF2
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        text = text.strip()
        if text:
            print(f"[ResumeParser] PDF parsed successfully with PyPDF2 ({len(text)} chars).")
            return text
        print("[ResumeParser] PyPDF2 returned empty — PDF appears to be image-based.")
    except ImportError:
        print("[ResumeParser] PyPDF2 not installed.")
    except Exception as e:
        print(f"[ResumeParser] PyPDF2 failed: {e}")

    # ── Vision OCR Fallback (Scanned PDF Detection) ──────────────────────────
    print("[ResumeParser] All text extractors returned empty. Treating as scanned PDF. Initiating Vision OCR...")
    try:
        import fitz
        from google import genai
        from PIL import Image
        import io
        from django.conf import settings
        
        # 1. Convert first 2 pages to low-quality JPEG (Memory Guarded)
        img_payloads = ["Extract all text from these resume images professionally. Output only the extracted text."]
        with fitz.open(path) as doc:
            for page_num in range(min(2, len(doc))):
                page = doc.load_page(page_num)
                mat = fitz.Matrix(150 / 72, 150 / 72)  # 150 DPI
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Further compress to 60% JPEG quality using Pillow to save RAM
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=60, optimize=True)
                img_bytes = buf.getvalue()
                
                img_payloads.append({'mime_type': 'image/jpeg', 'data': img_bytes})
                
                del page
                del pix
                del img
                buf.close()
            
        import gc; gc.collect()
        
        # 2. Gemini Vision Call
        api_key = getattr(settings, 'GEMINI_API_KEY', '').strip()
        if api_key and api_key != 'your_gemini_api_key_here':
            client = genai.Client(api_key=api_key)
            
            response = client.models.generate_content(
                model='gemini-1.5-flash',
                contents=img_payloads
            )
            text = response.text.strip()
            if text:
                print(f"[ResumeParser] Vision OCR successful ({len(text)} chars).")
                return text
        else:
            print("[ResumeParser] Gemini API key missing for Vision OCR fallback.")
            
    except Exception as e:
        if "MemoryError" in str(type(e)) or "Memory" in str(e) or "OOM" in str(e) or "cannot allocate" in str(e).lower():
            raise Exception("This image PDF is too large for our beta server. Please use a text-based PDF.")
        print(f"[ResumeParser] Vision OCR failed: {e}")

    return ""


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
        print(f"[ResumeParser] DOCX parsed successfully ({len(text)} chars).")
        return text
    except Exception as e:
        print(f"[ResumeParser] DOCX error: {e}")
        return ""


# ---------------------------------------------------------------------------
# Field extractors
# ---------------------------------------------------------------------------
def extract_email(text: str) -> str:
    """Return the first email address found in the text."""
    pattern = r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    return match.group() if match else ""


def sanitise_extracted_phone(phone: str) -> str:
    """
    Validate that a candidate phone string is not actually a date range or year.

    Rejection rules:
      R1 – Year-range pattern  : YYYY-YYYY or YYYY–YYYY  (e.g. 2020-2022)
      R2 – Standalone year     : any 4-digit value 1900-2099
      R3 – Too few digits      : fewer than 7 numeric digits
      R4 – Contains letters    : real phone numbers do not contain a-z

    Returns the original phone if it passes all checks, else ''.
    """
    if not phone:
        return ''

    # R1: year-range pattern
    if re.search(r'\b(19|20)\d{2}\s*[-\u2013\u2014]\s*(19|20)\d{2}\b', phone):
        print(f'[PhoneGuard] Rejected year-range: {phone!r}')
        return ''

    # R2: standalone 4-digit calendar year
    if re.search(r'\b(19|20)\d{2}\b', phone):
        print(f'[PhoneGuard] Rejected year-containing value: {phone!r}')
        return ''

    # R3: digit count
    digits_only = re.sub(r'\D', '', phone)
    if len(digits_only) < 7:
        print(f'[PhoneGuard] Rejected too-short value: {phone!r}')
        return ''

    # R4: alphabetic characters (dates like "Jan 2021" would be caught by R2,
    #     but catch any remaining word characters just in case)
    if re.search(r'[a-zA-Z]', phone):
        print(f'[PhoneGuard] Rejected alpha-containing value: {phone!r}')
        return ''

    return phone


def extract_phone(text: str) -> str:
    """
    Return the first valid phone number found in the text.

    Strategy:
      1. Pre-screen each line — skip lines that look like date ranges.
      2. Apply a phone-specific regex (requires 7-15 digits, standard separators).
      3. Post-validate with sanitise_extracted_phone() before returning.
    """
    # Lines that look like year ranges or education/experience dates are skipped
    _DATE_LINE = re.compile(
        r'\b(19|20)\d{2}\s*[-\u2013\u2014]\s*((19|20)\d{2}|present|current|now|till)\b',
        re.IGNORECASE,
    )

    # Strict phone pattern: starts with optional +, 7–15 digits, standard separators only
    _PHONE_RE = re.compile(
        r'(?<![0-9])(\+?[\d]{1,3}[\s\-.])?'   # optional country code
        r'(\(?\d{2,4}\)?[\s\-.])'              # area code
        r'(\d{3,4}[\s\-.])'                    # exchange
        r'(\d{3,4})'                            # subscriber
        r'(?![0-9])'
    )

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        # Skip lines that are clearly date ranges
        if _DATE_LINE.search(stripped):
            continue
        m = _PHONE_RE.search(stripped)
        if m:
            candidate = re.sub(r'\s+', ' ', m.group()).strip()
            validated = sanitise_extracted_phone(candidate)
            if validated:
                return validated[:20]

    return ''


def extract_name(text: str) -> str:
    """
    Heuristic: the candidate name is usually on the first non-empty line
    that is not an email, URL, or phone number.
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    skip_patterns = re.compile(
        r'(@|http|www\.|linkedin|github|resume|curriculum|vitae|\d{7,})',
        re.IGNORECASE
    )
    for line in lines[:6]:
        if not skip_patterns.search(line) and len(line.split()) <= 5:
            # Looks like a name (short, no special markers)
            return line.title()
    return lines[0].title() if lines else "Candidate"


def extract_sections(text: str) -> dict:
    """
    Split the resume text into labelled sections such as
    Education, Experience, Skills, Projects, etc.
    Returns a dict {section_name: section_text}.
    """
    section_headers = [
        'education', 'experience', 'work experience', 'skills',
        'technical skills', 'projects', 'certifications', 'achievements',
        'summary', 'objective', 'publications', 'languages', 'hobbies',
        'internship', 'training', 'volunteer', 'awards',
    ]

    # Build regex that matches any header (case-insensitive, on its own line)
    header_regex = re.compile(
        r'^\s*(' + '|'.join(re.escape(h) for h in section_headers) + r')\s*[:\-]?\s*$',
        re.IGNORECASE | re.MULTILINE
    )

    sections = {}
    lines = text.split('\n')
    current_section = 'header'
    current_lines = []

    for line in lines:
        if header_regex.match(line):
            sections[current_section] = '\n'.join(current_lines).strip()
            current_section = line.strip().lower().rstrip(':').strip()
            current_lines = []
        else:
            current_lines.append(line)

    sections[current_section] = '\n'.join(current_lines).strip()
    return sections
